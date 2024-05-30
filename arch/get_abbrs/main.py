import fitz  # Импортируем библиотеку PyMuPDF
import re
import os
import json
import logging

from docx import Document

from dict import abbrs
from pdf_searcher import search_pdf
from tex_strs import intro_strs, outro_strs

def get_abbrs(word_list):
    # оставляем только слова по шаблону - первые две бкувы заглавные - остальные любые
    new_word_list = []
    for word in word_list:
        cleaned_string = re.sub(r'^[\(]', '', word)
        cleaned_string = re.sub(r'[\)\»]*$', '', cleaned_string)
        cleaned_string = re.sub(r'\d+$', '', cleaned_string)
        if re.match('^[A-ZА-Я]{2}[A-Za-zА-Яа-я]*$', cleaned_string):
            new_word_list.append(cleaned_string)
    #print(new_word_list)
    abbrs = []
    for word in new_word_list:
        if len(word)<=7:
            abbrs.append(word)
    set_abbrs = set(abbrs)
    return list(set_abbrs)

def extract_words_from_pdf(pdf_path):
    words = []  # Создаем пустой список для слов
    inside_toa = False
    doc = fitz.open(pdf_path)  # Открываем PDF файл
    for page_number in range(doc.page_count):  # Перебираем все страницы
        page = doc.load_page(page_number)  # Загружаем страницу
        text = page.get_text("text")  # Получаем текст со страницы
        if 'Перечень принятых сокращений' in text:
            inside_toa = True
            continue
        if inside_toa and 'Устройство и работа' in text:
            inside_toa = False
            words += text.split()  # Добавляем слова в список
            continue
        if inside_toa and not 'Устройство и работа' in text:
            continue
        words += text.split()  # Добавляем слова в список
    return words  # Возвращаем список слов

# Configure logging
logging.basicConfig(filename='abbrs.log', level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logging.info("Запуск скрипта...")
# Ищем файл со словарем
if os.path.isfile('dictionary.json'):
    with open('dictionary.json', 'r', encoding='utf-8') as file:
        data = json.load(file)
        #print('found dictionary.json')
        logging.info("Найден внешний словарь абревиатур dictionary.json")
else:
    # Использование словаря "abbrs" в случае отсутствия файла "dict.json"
    data = abbrs
    logging.warning("Не найден внешний словарь абревиатур dictionary.json, будет использоваться пустой внутренний словарь!")

paths_to_pdfs = search_pdf()

if not paths_to_pdfs:
    logging.error("Нет файлов pdf в текущей папке для обработки...")

for path_to_pdf in paths_to_pdfs:
    logging.info(f"Обработка {path_to_pdf[0]}")
    word_list = extract_words_from_pdf(path_to_pdf[0])

    # убираем повторяющиеся слова
    word_set = set(word_list)
    word_list = sorted(list(word_set))

    # вытаскиваем абревиатуры
    new_word_list = get_abbrs(word_list)
    #print(word_list)
    #print(new_word_list)  # Выводим список всех слов

    new_word_list = sorted(new_word_list)
    ####################################################
    # выведем список абревиатур в файл !!! для сбора словаря - потом можно отключить
    # Имя файла, в который нужно сохранить список строк
    # file_path = "abbrs.txt"
    with open(path_to_pdf[2], 'w', encoding='utf-8') as file:
        for line in new_word_list:
            file.write(line + ', ')
    ####################################################

    used_keys = []
    tex_list = []
    doc_list = []

for word in new_word_list:
    # Проверяем, встречается ли ключ словаря в списке слов и не использовался ли уже
    if word in data.keys() and word not in used_keys:
        used_keys.append(word)
        value = data[word]
        #value = value[0].lower() + value[1:] # с маленькой буквы ?
        # Формируем строку tex и добавляем ее в tex_list
        tex_list.append(f'{word} & -- & {value}; \\\\'+'\n')
        doc_list.append(f'{word} - {value}')

# Меняем в последней строке ; на точку
if tex_list:
        last_element_index = len(tex_list) - 1
        last_element = tex_list[last_element_index]
        updated_last_element = last_element.replace('; \\\\\n', '. \\\\\n')
        tex_list[last_element_index] = updated_last_element   

'''
    for word in new_word_list:
        # Проверка вхождения ключа словаря в строку
        for key, value in data.items():
            if key == word and key not in used_keys:
                #print(f"Ключ '{key}' найден в строке, соответствующее значение: '{value}'.")
                used_keys.append(key)
                tex_list.append(f'{key} & -- & {value} \\\\'+'\n')
'''
final_tex = intro_strs + tex_list + outro_strs

# Открываем файл для записи в UTF-8
with open(path_to_pdf[1], 'w', encoding='utf-8') as file:
    for line in final_tex:
        file.write(line)  # Добавляем символ новой строки после каждой строки

logging.info("Создаем перечень аббревиатур в формате doc...")

doc = Document()
# Добавление точки с запятой в конец каждой строки
for i in range(len(doc_list)-1):
    doc_list[i] += ';'

# Добавление точки в последнюю строку
doc_list[-1] += '.'
doc.add_paragraph('Перечень принятых сокращений')
for item in doc_list:
    doc.add_paragraph(item)

# Сохранение документа
doc.save(path_to_pdf[3])

logging.info("Останов скрипта...")